"""Render a saved case as a Word document a lawyer can carry to the lectern."""

import io
import re

from docx import Document
from docx.shared import Pt

from app.models.schemas import Argument, CaseAnalysis, CitedFact

DISCLAIMER = (
    "AI-GENERATED — verify every fact against the cited source documents and every "
    "authority against the linked judgment before relying on this in court."
)


def _format_citations(fact: CitedFact) -> str:
    return "; ".join(f"{c.source_document}, {c.location}" for c in fact.citations)


def _add_fact(document: Document, fact: CitedFact) -> None:
    paragraph = document.add_paragraph(fact.text, style="List Bullet")
    citations = _format_citations(fact)
    if citations:
        run = paragraph.add_run(f"  [{citations}]")
        run.italic = True
        run.font.size = Pt(9)


def _add_argument(document: Document, index: int, argument: Argument) -> None:
    document.add_heading(f"{index}. {argument.point}", level=2)

    if argument.legal_basis:
        document.add_paragraph(f"Legal basis: {argument.legal_basis}")

    if argument.supporting_facts:
        document.add_paragraph("Supporting facts:")
        for fact in argument.supporting_facts:
            _add_fact(document, fact)

    if argument.authorities:
        document.add_paragraph("Authorities:")
        for authority in argument.authorities:
            parts = [authority.title]
            if authority.court:
                parts.append(authority.court)
            if authority.date:
                parts.append(authority.date)
            line = " — ".join(parts)
            paragraph = document.add_paragraph(line, style="List Bullet")
            run = paragraph.add_run(f"  {authority.url}")
            run.italic = True
            run.font.size = Pt(9)

    if argument.counter_argument:
        document.add_paragraph(f"Anticipated counter-argument: {argument.counter_argument}")
    if argument.rebuttal:
        document.add_paragraph(f"Rebuttal: {argument.rebuttal}")


def build_hearing_pack(title: str, analysis: CaseAnalysis) -> bytes:
    document = Document()

    document.add_heading(title, level=0)
    warning = document.add_paragraph(DISCLAIMER)
    warning.runs[0].italic = True

    prep = analysis.hearing_prep

    document.add_heading("Hearing Brief", level=1)
    document.add_paragraph(prep.brief if prep else "The Prepare pack could not be generated.")

    if prep and prep.outline:
        document.add_heading("Oral Argument Outline", level=1)
        for section in prep.outline:
            document.add_paragraph(section.heading, style="List Number")
            for point in section.talking_points:
                document.add_paragraph(point, style="List Bullet 2")

    document.add_heading("Arguments", level=1)
    if not analysis.arguments:
        document.add_paragraph("No arguments were generated.")
    for index, argument in enumerate(analysis.arguments, start=1):
        _add_argument(document, index, argument)

    if prep and prep.checklist:
        document.add_heading("Checklist", level=1)
        for entry in prep.checklist:
            document.add_paragraph(f"☐ ({entry.category}) {entry.item}")

    if analysis.warnings:
        document.add_heading("Warnings from this run", level=1)
        for note in analysis.warnings:
            document.add_paragraph(note, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_filename(title: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
    return f"{slug or 'hearing-pack'}.docx"
